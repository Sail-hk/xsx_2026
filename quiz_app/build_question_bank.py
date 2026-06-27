import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
QUESTION_BANK_DIR = ROOT_DIR / "题库"
OUTPUT_PATH = BASE_DIR / "questions-data.js"

QUESTION_RE = re.compile(r"^(\d+)[.\uFF0E\u3001]\s*(.+)$")
OPTION_RE = re.compile(r"^([A-D])[.\uFF0E\u3001]\s*(.+)$")
OPTION_INLINE_RE = re.compile(r"([A-D])[.\uFF0E\u3001]\s*(.*?)(?=\s+[A-D][.\uFF0E\u3001]\s*|$)")
INLINE_ANSWER_RE = re.compile(r"^(.*?)([A-D])$")
PAIR_RE = re.compile(r"(\d+)\s*[.\uFF0E]?\s*[:：]?\s*(A|B|C|D|正确|错误|ABCD|ABC|ABD|ACD|BCD|AC|AD|BC|BD|CD)")
CHAPTER_HEADING_RE = re.compile(r"^(第[一二三四五六七八九十百0-9]+章)")

GROUPED_CHAPTERS = {
    "第二三章",
    "第四五六章",
    "第七到十章",
    "第十一到十四章",
    "第十五到十七章",
}

CHAPTER_ORDER = [
    "导论",
    "第一章",
    "第二章",
    "第三章",
    "第四章",
    "第五章",
    "第六章",
    "第七章",
    "第八章",
    "第九章",
    "第十章",
    "第十一章",
    "第十二章",
    "第十三章",
    "第十四章",
    "第十五章",
    "第十六章",
    "第十七章",
    "综合练习一",
    "综合练习二",
]

COMPREHENSIVE_LABELS = {
    "综合知识点练习题_无章节提示版(1).docx": "综合练习一",
    "综合知识点练习题_第二套_无章节提示版.docx": "综合练习二",
}

SUPPLEMENT_QUESTIONS = [
    {
        "type": "multiple",
        "chapter": "导论",
        "stem": "习近平新时代中国特色社会主义思想的主要内容，可以概括为哪些方面？",
        "options": [
            {"key": "A", "text": "十个明确"},
            {"key": "B", "text": "十四个坚持"},
            {"key": "C", "text": "十三个方面成就"},
            {"key": "D", "text": "六个必须坚持"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第一章",
        "stem": "中国式现代化的鲜明特征包括哪些？",
        "options": [
            {"key": "A", "text": "人口规模巨大的现代化"},
            {"key": "B", "text": "全体人民共同富裕的现代化"},
            {"key": "C", "text": "人与自然和谐共生的现代化"},
            {"key": "D", "text": "走和平发展道路的现代化"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第一章",
        "stem": "中国式现代化的本质要求包括哪些？",
        "options": [
            {"key": "A", "text": "坚持中国共产党领导"},
            {"key": "B", "text": "实现高质量发展"},
            {"key": "C", "text": "推动构建人类命运共同体"},
            {"key": "D", "text": "照搬西方现代化模式"},
        ],
        "answer": "ABC",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第四五六章",
        "stem": "坚持人民立场的基本要求包括哪些？",
        "options": [
            {"key": "A", "text": "始终牢记党的初心和使命"},
            {"key": "B", "text": "始终保持党同人民群众的血肉联系"},
            {"key": "C", "text": "热爱人民、尊重人民、敬畏人民"},
            {"key": "D", "text": "以资本增值作为根本目标"},
        ],
        "answer": "ABC",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第四五六章",
        "stem": "全面深化改革总目标的核心表述包括哪些？",
        "options": [
            {"key": "A", "text": "完善和发展中国特色社会主义制度"},
            {"key": "B", "text": "推进国家治理体系现代化"},
            {"key": "C", "text": "推进治理能力现代化"},
            {"key": "D", "text": "全面取消政府调控作用"},
        ],
        "answer": "ABC",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第七到十章",
        "stem": "新发展理念包括哪些内容？",
        "options": [
            {"key": "A", "text": "创新"},
            {"key": "B", "text": "协调"},
            {"key": "C", "text": "绿色"},
            {"key": "D", "text": "共享"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第七到十章",
        "stem": "公有制经济包括哪些成分？",
        "options": [
            {"key": "A", "text": "国有经济"},
            {"key": "B", "text": "集体经济"},
            {"key": "C", "text": "混合所有制经济中的国有成分"},
            {"key": "D", "text": "混合所有制经济中的集体成分"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第七到十章",
        "stem": "教育、科技、人才被定位为什么？",
        "options": [
            {"key": "A", "text": "全面建设社会主义现代化国家的基础性支撑"},
            {"key": "B", "text": "全面建设社会主义现代化国家的战略性支撑"},
            {"key": "C", "text": "推动高质量发展的重要力量"},
            {"key": "D", "text": "可有可无的辅助条件"},
        ],
        "answer": "ABC",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第七到十章",
        "stem": "全过程人民民主的相关表述，正确的有哪些？",
        "options": [
            {"key": "A", "text": "人民民主是社会主义的生命"},
            {"key": "B", "text": "全过程人民民主是社会主义民主政治的本质属性"},
            {"key": "C", "text": "它建立在社会主义经济基础之上"},
            {"key": "D", "text": "它是一种脱离制度安排的抽象民主"},
        ],
        "answer": "ABC",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第十一到十四章",
        "stem": "中华文明的突出特性包括哪些？",
        "options": [
            {"key": "A", "text": "连续性"},
            {"key": "B", "text": "创新性"},
            {"key": "C", "text": "统一性"},
            {"key": "D", "text": "包容性"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第十一到十四章",
        "stem": "总体国家安全观强调哪些原则？",
        "options": [
            {"key": "A", "text": "以人民安全为宗旨"},
            {"key": "B", "text": "以政治安全为根本"},
            {"key": "C", "text": "以经济安全为基础"},
            {
                "key": "D",
                "text": "维护国家安全要贯穿党和国家工作各方面全过程",
            },
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第十五到十七章",
        "stem": "新时代党的强军目标包括哪些内容？",
        "options": [
            {"key": "A", "text": "听党指挥"},
            {"key": "B", "text": "能打胜仗"},
            {"key": "C", "text": "作风优良"},
            {"key": "D", "text": "把人民军队建设成为世界一流军队"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "multiple",
        "chapter": "第十五到十七章",
        "stem": "全面从严治党的相关表述，正确的有哪些？",
        "options": [
            {"key": "A", "text": "党的政治建设是党的根本性建设"},
            {"key": "B", "text": "思想建设是基础性建设"},
            {"key": "C", "text": "作风问题核心是党同人民群众的关系问题"},
            {"key": "D", "text": "反腐倡廉的核心是制度制约和监督权力"},
        ],
        "answer": "ABCD",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "导论",
        "stem": "2017年10月党的十九大把习近平新时代中国特色社会主义思想写入党章，2018年3月又载入宪法。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "导论",
        "stem": "文化自信是更基础、更广泛、更深厚的自信。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第一章",
        "stem": "新时代我国社会主要矛盾发生变化，说明我国已经不再处于社会主义初级阶段。",
        "options": [],
        "answer": "错误",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第四五六章",
        "stem": "改革开放是决定当代中国命运的关键一招。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第四五六章",
        "stem": "国家治理体系和治理能力，是一个国家制度和制度执行能力的集中体现。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第七到十章",
        "stem": "高质量发展是全面建设社会主义现代化国家的首要任务。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第七到十章",
        "stem": "教育、科技、人才只是经济工作的辅助条件，不属于现代化建设的战略性支撑。",
        "options": [],
        "answer": "错误",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第七到十章",
        "stem": "公正司法是维护社会公平正义的最后一道防线。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第十一到十四章",
        "stem": "良好生态环境是最公平的公共产品，是最普惠的民生福祉。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第十一到十四章",
        "stem": "政治安全是国家安全的根本。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第十五到十七章",
        "stem": "军委主席负责制是坚持党对人民军队绝对领导的根本制度和根本实现形式。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第十五到十七章",
        "stem": "多边主义不符合各国人民利益，不是维护和平、促进发展的有效路径。",
        "options": [],
        "answer": "错误",
        "sources": ["习思想重点.docx"],
    },
    {
        "type": "judge",
        "chapter": "第十五到十七章",
        "stem": "党的领导是中国特色社会主义最本质的特征，也是中国特色社会主义制度的最大优势。",
        "options": [],
        "answer": "正确",
        "sources": ["习思想重点.docx"],
    },
]


def iter_docx_files() -> list[Path]:
    return sorted(
        path
        for path in QUESTION_BANK_DIR.glob("*.docx")
        if path.is_file() and not path.name.startswith("~$")
    )


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "").strip()


def normalize_answer(answer: str, question_type: str) -> str:
    normalized = normalize_text(answer).upper()
    if question_type == "judge":
        if normalized in {"正确", "T", "TRUE"}:
            return "正确"
        if normalized in {"错误", "F", "FALSE"}:
            return "错误"
    return "".join(sorted(normalized))


def clean_stem_with_answer(stem: str, answer: str, question_type: str) -> str:
    cleaned_stem = stem.strip()
    normalized_answer = normalize_answer(answer, question_type)

    if question_type == "single" and normalized_answer in "ABCD":
        patterns = [
            rf"^(.*?)[\s\u3000]*[（\(]?\s*{re.escape(normalized_answer)}\s*[）\)]?$",
            rf"^(.*?[。！？?])[\s\u3000]*{re.escape(normalized_answer)}$",
        ]
        for pattern in patterns:
            match = re.match(pattern, cleaned_stem)
            if match:
                candidate = match.group(1).strip()
                if len(candidate) >= 4:
                    cleaned_stem = candidate
                    break

    if question_type == "judge":
        patterns = [
            r"^(.*?)[\s\u3000]*[（\(]?\s*(正确|错误)\s*[）\)]?$",
            r"^(.*?[。！？?])[\s\u3000]*(正确|错误)$",
        ]
        for pattern in patterns:
            match = re.match(pattern, cleaned_stem)
            if match:
                candidate = match.group(1).strip()
                if len(candidate) >= 4:
                    cleaned_stem = candidate
                    break

    return cleaned_stem


def infer_chapter(filename: str) -> str:
    if filename in COMPREHENSIVE_LABELS:
        return COMPREHENSIVE_LABELS[filename]
    if filename.startswith("1."):
        return "导论"
    if filename.startswith("2."):
        return "第一章"
    if filename.startswith("3."):
        return "第二三章"
    if filename.startswith("4."):
        return "第四五六章"
    if filename.startswith("5."):
        return "第七到十章"
    if filename.startswith("6."):
        return "第十一到十四章"
    if filename.startswith("7."):
        return "第十五到十七章"
    return "综合练习一"


def parse_answer_tables(doc: Document) -> dict[int, str]:
    answers: dict[int, str] = {}
    for table in doc.tables:
        for row in table.rows:
            pending_number = None
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                for number, answer in PAIR_RE.findall(text):
                    answers[int(number)] = answer
                if pending_number is not None and re.fullmatch(
                    r"(A|B|C|D|正确|错误|ABCD|ABC|ABD|ACD|BCD|AC|AD|BC|BD|CD)",
                    text,
                ):
                    answers[pending_number] = text
                    pending_number = None
                    continue
                if text.isdigit():
                    pending_number = int(text)
    return answers


def parse_answer_paragraphs(lines: list[str]) -> dict[int, str]:
    answers: dict[int, str] = {}
    in_answer_block = False
    for line in lines:
        if line == "参考答案":
            in_answer_block = True
            continue
        if not in_answer_block:
            continue
        for number, answer in PAIR_RE.findall(line):
            answers[int(number)] = answer
    return answers


def is_section_heading(line: str) -> bool:
    return "单项选择题" in line or "多项选择题" in line or "判断题" in line


def infer_question_type(line: str, current_type: str) -> str:
    if "多项选择题" in line:
        return "multiple"
    if "判断题" in line:
        return "judge"
    if "单项选择题" in line:
        return "single"
    return current_type


def should_skip_line(line: str) -> bool:
    if line == "参考答案":
        return True
    if line.startswith("说明：") or line.startswith("题型："):
        return True
    if re.fullmatch(r"第[一二三四五六七八九十百]+章", line):
        return True
    if re.fullmatch(r"[一二三四五六七八九十]+、", line) and not is_section_heading(line):
        return True
    return False


def infer_chapter_from_heading(line: str, default_chapter: str) -> str | None:
    if default_chapter not in GROUPED_CHAPTERS:
        return None
    match = CHAPTER_HEADING_RE.match(line)
    if not match:
        return None
    return match.group(1)


def parse_questions_from_doc(path: Path) -> list[dict]:
    doc = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_answers = parse_answer_tables(doc)
    paragraph_answers = parse_answer_paragraphs(lines)
    answer_map = {**paragraph_answers, **table_answers}

    questions: list[dict] = []
    current_question = None
    current_type = "single"
    inline_answer_mode = len(doc.tables) == 0
    default_chapter = infer_chapter(path.name)
    current_chapter = default_chapter

    def flush_current() -> None:
        nonlocal current_question
        if current_question is None:
            return
        answer = current_question.get("answer") or answer_map.get(current_question["number"])
        if not answer:
            current_question = None
            return
        current_question["stem"] = clean_stem_with_answer(
            current_question["stem"], answer, current_question["type"]
        )
        current_question["answer"] = normalize_answer(answer, current_question["type"])
        current_question["source"] = path.name
        current_question["chapter"] = current_chapter
        questions.append(current_question)
        current_question = None

    for line in lines:
        if line == "参考答案":
            break

        heading_chapter = infer_chapter_from_heading(line, default_chapter)
        if heading_chapter:
            flush_current()
            current_chapter = heading_chapter
            continue

        if is_section_heading(line):
            flush_current()
            current_type = infer_question_type(line, current_type)
            continue

        if should_skip_line(line):
            continue

        question_match = QUESTION_RE.match(line)
        if question_match:
            flush_current()
            number = int(question_match.group(1))
            stem = question_match.group(2).strip()
            answer = None

            if inline_answer_mode and current_type != "judge":
                inline_match = INLINE_ANSWER_RE.match(stem)
                if inline_match and inline_match.group(2) in "ABCD":
                    stem = inline_match.group(1).rstrip("（(。！？?")
                    answer = inline_match.group(2)

            current_question = {
                "number": number,
                "stem": stem,
                "type": current_type,
                "options": [],
                "answer": answer,
            }
            continue

        if current_question is None:
            continue

        option_match = OPTION_RE.match(line)
        if option_match:
            current_question["options"].extend(
                {"key": match.group(1), "text": match.group(2).strip()}
                for match in OPTION_INLINE_RE.finditer(line)
            )

    flush_current()
    return questions


def resolve_grouped_supplement_chapter(question: dict) -> str:
    stem = question["stem"]

    if "坚持人民立场" in stem:
        return "第四章"
    if "全面深化改革总目标" in stem or "国家治理体系和治理能力" in stem or "改革开放是决定当代中国命运的关键一招" in stem:
        return "第五章"
    if "新发展理念" in stem or "公有制经济" in stem or "高质量发展" in stem:
        return "第六章"
    if "教育、科技、人才" in stem:
        return "第七章"
    if "全过程人民民主" in stem:
        return "第八章"
    if "公正司法" in stem:
        return "第九章"
    if "中华文明的突出特性" in stem:
        return "第十章"
    if "良好生态环境" in stem:
        return "第十二章"
    if "总体国家安全观" in stem or "政治安全是国家安全的根本" in stem:
        return "第十三章"
    if "新时代党的强军目标" in stem or "军委主席负责制" in stem:
        return "第十四章"
    if "多边主义" in stem:
        return "第十六章"
    if "全面从严治党" in stem or "党的领导是中国特色社会主义最本质的特征" in stem:
        return "第十七章"
    return question["chapter"]


def refine_question_chapter(question: dict) -> dict:
    refined = dict(question)
    if refined.get("chapter") in GROUPED_CHAPTERS:
        refined["chapter"] = resolve_grouped_supplement_chapter(refined)
    return refined


def deduplicate_questions(questions: list[dict]) -> list[dict]:
    deduped = []
    seen = {}
    for raw_question in questions:
        question = refine_question_chapter(raw_question)
        sources = question.get("sources") or [question["source"]]
        key = (
            question["type"],
            question.get("chapter", "综合练习一"),
            normalize_text(question["stem"]),
            tuple((item["key"], normalize_text(item["text"])) for item in question["options"]),
        )
        if key in seen:
            seen[key]["sources"].extend(sources)
            continue
        entry = {
            "id": f"{question['type']}-{len(deduped) + 1}",
            "type": question["type"],
            "chapter": question.get("chapter", "综合练习一"),
            "stem": question["stem"],
            "options": question["options"],
            "answer": question["answer"],
            "sources": list(sources),
        }
        deduped.append(entry)
        seen[key] = entry
    return deduped


def build_question_bank() -> dict:
    source_files = list(iter_docx_files())
    all_questions = []
    for docx_path in source_files:
        all_questions.extend(parse_questions_from_doc(docx_path))
    all_questions.extend(SUPPLEMENT_QUESTIONS)

    questions = deduplicate_questions(all_questions)
    by_type = defaultdict(list)
    by_chapter = defaultdict(list)

    for question in questions:
        by_type[question["type"]].append(question)
        by_chapter[question["chapter"]].append(question)

    ordered_chapters = {
        chapter: by_chapter[chapter]
        for chapter in CHAPTER_ORDER
        if by_chapter[chapter]
    }

    return {
        "meta": {
            "sourceDir": QUESTION_BANK_DIR.name,
            "sourceFiles": [path.name for path in source_files],
            "counts": {key: len(value) for key, value in by_type.items()},
            "chapterCounts": {key: len(value) for key, value in ordered_chapters.items()},
            "chapterOrder": list(ordered_chapters.keys()),
            "total": len(questions),
            "supplementCount": len(SUPPLEMENT_QUESTIONS),
        },
        "questions": {
            "single": by_type["single"],
            "multiple": by_type["multiple"],
            "judge": by_type["judge"],
        },
        "chapters": ordered_chapters,
    }


def main() -> None:
    bank = build_question_bank()
    payload = "window.QUESTION_BANK = " + json.dumps(bank, ensure_ascii=False, indent=2) + ";\n"
    OUTPUT_PATH.write_text(payload, encoding="utf-8")
    print(f"Question bank written to {OUTPUT_PATH}")
    print(json.dumps(bank["meta"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
